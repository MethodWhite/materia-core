#!/usr/bin/env python3
"""
Genera ambos documentos DOCX:
1. HSAQ_STANDARD.docx — Estándar completo de HyperSparse Adaptive Quantization
2. MATERIA_ARQUITECTURA.docx — Arquitectura Toroidal Hexagonal M.A.T.E.R.I.A. V4
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENTATION

OUT_DIR = "/home/methodwhite/Documentos/IA LATAM"

# ─── Helper utilities ───────────────────────────────────────

def setup_page(doc):
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)  # A4
    sec.top_margin = sec.bottom_margin = Cm(2.54)
    sec.left_margin = sec.right_margin = Cm(3.18)

def tune_styles(doc):
    body = doc.styles["Normal"]
    body.font.name = "Calibri"
    body.font.size = Pt(11)
    body.paragraph_format.line_spacing = 1.15
    body.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri Light"
        h.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        h.paragraph_format.space_before = Pt(12 if level == 1 else 8)
        h.paragraph_format.space_after = Pt(6)
    doc.styles["Heading 1"].font.size = Pt(18)
    doc.styles["Heading 2"].font.size = Pt(14)
    doc.styles["Heading 3"].font.size = Pt(12)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text, bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def add_code(doc, code):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.paragraph_format.space_after = Pt(2)
    for line in code.split("\n"):
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    return p

def make_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = "Calibri"
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Header shading
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd",
                                  {"{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill": "1F3A5F",
                                   "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color": "auto",
                                   "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val": "clear"})
        shading.append(shd)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.text = val
            for r in cell.paragraphs[0].runs:
                r.font.name = "Calibri"
                r.font.size = Pt(10)
    return t

def add_caption(doc, text):
    p = doc.add_paragraph(text, style="Caption")
    for r in p.runs:
        r.font.size = Pt(9)
        r.italic = True
        r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

# ═══════════════════════════════════════════════════════════════
# DOCUMENTO 1: HSAQ STANDARD
# ═══════════════════════════════════════════════════════════════

def generate_hsaq_docx():
    path = os.path.join(OUT_DIR, "HSAQ", "HSAQ_STANDARD.docx")
    os.makedirs(os.path.dirname(path), exist_ok=True)

    doc = Document()
    setup_page(doc)
    tune_styles(doc)

    # ── Portada ──
    for _ in range(6):
        doc.add_paragraph()
    add_para(doc, "HSAQ Standard", bold=True, size=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "HyperSparse Adaptive Quantization", size=18, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    add_para(doc, "Versión 1.1 — Julio 2026", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para(doc, "M.A.T.E.R.I.A. Research", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Congreso AI LATAM 2026", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    doc.add_page_break()

    # ── 1. Definición ──
    add_heading(doc, "1. Definición", 1)
    add_para(doc, (
        "HSAQ (HyperSparse Adaptive Quantization) es un mecanismo de cuantización de "
        "activaciones mediante sparsity adaptativa dinámica, donde cada capa de la red "
        "calcula su propio umbral vía torch.kthvalue en cada forward pass."
    ))
    add_para(doc, (
        'La "cuantización" en HSAQ refiere a que las activaciones se reducen a un conjunto '
        "discreto de valores {0, valor_original} mediante una máscara binaria calculada "
        "dinámicamente. NO es cuantización de pesos. NO es INT8/INT4. NO es bitsandbytes."
    ), bold=True)

    # ── 2. Principios ──
    add_heading(doc, "2. Principios Fundamentales", 1)
    principles = [
        "Adaptativo: el umbral de sparsity se recalcula en cada batch usando kthvalue",
        "Por capa: cada componente del modelo tiene su propio umbral independiente",
        "Sin calibración: no requiere datasets externos ni pasos post-entrenamiento",
        "Hardware-agnostic: funciona en CPU, GPU y TPU sin modificaciones",
        "Sin estado: no hay buffers persistentes entre batches",
        "Gradiente fluye: las neuronas activas reciben gradiente normalmente (STE nativo)",
    ]
    for p in principles:
        add_bullet(doc, p)

    # ── 3. Algoritmo ──
    add_heading(doc, "3. Algoritmo", 1)
    add_heading(doc, "3.1 Pseudocódigo", 2)
    add_code(doc, """Entrada: x ∈ ℝ^(B×T×D)      Batch de activaciones
Parámetro: s ∈ [0,1)          Fracción a enmascarar

1. magnitudes = |x|.view(B, -1)        Magnitudes por batch
2. k = max(1, n · s)                  Neuronas a enmascarar
3. umbral = kthvalue(magnitudes, k)   k-ésimo valor más pequeño
4. máscara = |x| ≥ umbral             Máscara binaria {0, 1}
5. return x × máscara                 Solo neuronas relevantes""")

    add_heading(doc, "3.2 Implementación de Referencia", 2)
    add_code(doc, """class HSAQ(nn.Module):
    def __init__(self, sparsity=0.3):
        super().__init__()
        self.sparsity = sparsity

    def forward(self, x, sparsity_override=None):
        s = sparsity_override or self.sparsity
        if s <= 0: return x
        flat = x.abs().view(x.size(0), -1)
        k = max(1, min(n-1, int(n * s)))
        thresh = torch.kthvalue(flat, k, dim=1).values
        thresh = thresh.view(-1, *([1]*(x.dim()-1)))
        return x * (x.abs() >= thresh)""")

    # ── 4. Puntos de Aplicación ──
    add_heading(doc, "4. Puntos de Aplicación", 1)
    make_table(doc,
        ["#", "Componente", "Sparsity", "Propósito"],
        [
            ["1", "Embedding", "5%", "Preservar información de entrada"],
            ["2", "Transformer t2", "8%", "Capas tempranas"],
            ["3", "Transformer t5", "12%", "Capas medias"],
            ["4", "Transformer t8", "15%", "Capas tardías"],
            ["5", "LIF-SNN", "10%", "Procesamiento temporal"],
            ["6", "SSM", "5%", "Estado latente"],
            ["7", "JEPA Hub", "5%", "Espacio de predicción"],
        ]
    )
    add_caption(doc, "Tabla 1: Puntos de aplicación HSAQ en M.A.T.E.R.I.A. V4")

    # ── 5. Sparsity Escalonada ──
    add_heading(doc, "5. Sparsity Escalonada", 1)
    add_para(doc, (
        "El problema crítico identificado fue el efecto compuesto (compounding sparsity): "
        "aplicar HSAQ con sparsity uniforme del 30% en múltiples puntos del pipeline resultaba "
        "en solo 0.7% de información sobreviviente. La solución:"
    ))
    add_bullet(doc, "Capas tempranas: sparsity baja (5-8%) para preservar información general")
    add_bullet(doc, "Capas medias: sparsity media (10-12%) para compresión gradual")
    add_bullet(doc, "Capas tardías: sparsity moderada (15%) para eliminar ruido")
    add_bullet(doc, "Componentes finales: sparsity baja (5%) para preservar representaciones")
    add_para(doc, (
        "Con sparsity escalonada y 7 puntos estratégicos × 2 ciclos toroidales, "
        "la información preservada es aproximadamente 52% (vs 0.7% con uniforme 30%)."
    ))

    # ── 6. Optimizer ──
    add_heading(doc, "6. HSAQ como Optimizer", 1)
    add_para(doc, (
        "HSAQ reemplaza a AdamW como mecanismo de optimización. La máscara sparse actúa "
        "como regularizador adaptativo: las neuronas irrelevantes no reciben gradiente."
    ))
    make_table(doc,
        ["Aspecto", "AdamW", "HSAQ + SGD Nesterov"],
        [
            ["Estados de optimizer", "2 por parámetro (8 bytes)", "1 por parámetro (4 bytes)"],
            ["Regularización", "weight_decay", "Sparsity adaptativa"],
            ["Memoria extra (190M params)", "~1.5 GB", "~0.76 GB"],
            ["Convergencia", "Rápida", "Comparable con +48% accuracy"],
        ]
    )
    add_caption(doc, "Tabla 2: Comparativa AdamW vs HSAQ + SGD Nesterov")

    # ── 7. Comparación ──
    add_heading(doc, "7. Comparación con TurboQuant", 1)
    make_table(doc,
        ["Característica", "TurboQuant (Google)", "HSAQ (MATERIA)"],
        [
            ["Tipo", "Cuantización INT8 de pesos", "Sparsity de activaciones"],
            ["Adaptativo", "No (fijo post-calibración)", "Sí (kthvalue por batch)"],
            ["Calibración", "Requiere dataset externo", "No requiere"],
            ["Hardware", "Solo GPU con INT8", "CPU/GPU/TPU"],
            ["Info. preservada", "~99%", "~52% (controlada)"],
        ]
    )
    add_caption(doc, "Tabla 3: HSAQ vs TurboQuant")

    # ── 8. Integración ──
    add_heading(doc, "8. Integración en M.A.T.E.R.I.A. V4", 1)
    add_para(doc, (
        "En M.A.T.E.R.I.A. V4, HSAQ se integra en el toroide hexagonal con JEPA como hub "
        "central. Cada arista del hexágono tiene HSAQ con sparsity calibrada:"
    ))
    add_code(doc, """Input → Emb → HSAQ(5%)
   → JEPA Hub (central)
     → Transformer → HSAQ(8-15%) → JEPA
     → SNN → HSAQ(10%) → JEPA
     → SSM → HSAQ(5%) → JEPA
   → JEPA integra → HSAQ(5%) → Head
   → 2 ciclos toroidales completos""")

    # ── 9. Archivos ──
    add_heading(doc, "9. Archivos del Sistema", 1)
    make_table(doc,
        ["Archivo", "Propósito"],
        [
            ["models/core/hsaq.py", "Implementación de HSAQ (~60 líneas)"],
            ["models/materia_v4.py", "Modelo completo con HSAQ por capa"],
            ["configs/V4_210M_BPE.yaml", "Configuración de entrenamiento"],
            ["scripts/train_v4_enhanced.py", "Training script con tracking HSAQ"],
            ["scripts/plot_hsaq.py", "Visualización de métricas HSAQ"],
        ]
    )

    # ── 10. Limitaciones ──
    add_heading(doc, "10. Limitaciones y Trabajo Futuro", 1)
    add_bullet(doc, "Sparsity fija: actualmente configurada por capa, no aprendida")
    add_bullet(doc, "Sin kernel CUDA sparse: la máscara binaria no ahorra FLOPs reales")
    add_bullet(doc, "Export ONNX: kthvalue no está en opset estándar ONNX")
    add_bullet(doc, "Synapsis solo en inferencia: no durante entrenamiento")

    # ── Cierre ──
    add_para(doc, "", space_after=12)
    add_para(doc, "— Fin del documento —", italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(path)
    print(f"✅ HSAQ DOCX: {path}")

# ═══════════════════════════════════════════════════════════════
# DOCUMENTO 2: MATERIA ARQUITECTURA
# ═══════════════════════════════════════════════════════════════

def generate_materia_docx():
    path = os.path.join(OUT_DIR, "M.A.T.E.R.I.A.", "MATERIA_ARQUITECTURA_TOROIDAL.docx")
    os.makedirs(os.path.dirname(path), exist_ok=True)

    doc = Document()
    setup_page(doc)
    tune_styles(doc)

    # ── Portada ──
    for _ in range(5):
        doc.add_paragraph()
    add_para(doc, "M.A.T.E.R.I.A. V4", bold=True, size=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Arquitectura Toroidal Hexagonal", size=20, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    doc.add_paragraph()
    add_para(doc, "JEPA como hub central del toroide de interconexión hexagonal", size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Inspirado en geometría sagrada: flower of life, toro, hexágono sagrado", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    doc.add_paragraph()
    doc.add_paragraph()
    add_para(doc, "MethodWhite — M.A.T.E.R.I.A. Research", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Versión 4.0 — Julio 2026", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    doc.add_page_break()

    # ── 1. Visión General ──
    add_heading(doc, "1. Visión General", 1)
    add_para(doc, (
        "M.A.T.E.R.I.A. V4 es un modelo de inteligencia artificial con una arquitectura "
        "toroidal hexagonal, donde JEPA (Joint Embedding Predictive Architecture) actúa como "
        "hub central del toroide. Todos los componentes del modelo —Transformer, SNN, SSM, "
        "Embedding, Head, Synapsis— se conectan exclusivamente al espacio latente JEPA "
        "mediante proyecciones hexagonales bidireccionales."
    ))

    add_heading(doc, "1.1 Diagrama Arquitectural", 2)
    add_code(doc, """                ┌─────────────┐
               ╱  Transformer  ╲
              │     ↕ ↕ ↕      │
     ┌───────┐│   ↔ JEPA ↔    │┌───────┐
     │  SSM  │←━━→   ↕   ←━━→││  SNN  │
     └───────┘│  ↔  Hub  ↔   │└───────┘
              │     ↕ ↕ ↕      │
               ╲              ╱
                └─────┬───────┘
                      ↕
                ┌─────┴───────┐
                │   Head/Emb   │ ← Toroidal
                └─────────────┘""")

    add_heading(doc, "1.2 Principios de Diseño", 2)
    add_bullet(doc, "JEPA-First: JEPA es el hub central, no un componente periférico")
    add_bullet(doc, "Toroidal: el flujo de información es cíclico, no lineal")
    add_bullet(doc, "Hexagonal: 6 conexiones por componente (geometría sagrada)")
    add_bullet(doc, "HSAQ en cada arista: sparsity adaptativa en todas las conexiones")
    add_bullet(doc, "Multi-paradigma: Transformer + SNN + SSM + JEPA convergiendo")

    # ── 2. JEPA Hub ──
    add_heading(doc, "2. JEPA Hub Central", 1)
    add_para(doc, (
        "JEPA es el corazón de la arquitectura. No es un componente más al final del pipeline, "
        "sino el espacio latente central donde convergen y desde donde se distribuyen todas las "
        "señales del modelo."
    ))
    add_heading(doc, "2.1 Codificación al Espacio Latente", 2)
    add_code(doc, """class JEPAEncoder(nn.Module):
    def __init__(self, latent_dim):
        super().__init__()
        self.proj = nn.Linear(latent_dim, latent_dim * 2)
        self.norm = nn.RMSNorm(latent_dim * 2)
        self.out = nn.Linear(latent_dim * 2, latent_dim)

    def forward(self, x):
        x = self.proj(x)
        x = F.silu(self.norm(x))
        return self.out(x)""")

    add_heading(doc, "2.2 Predicción con SCA", 2)
    add_para(doc, (
        "El predictor JEPA utiliza descomposición espectral SCA con constante de acoplamiento "
        "K = √(π·e·γ) = 2.781042. Los autovalores λ_n = K · σ(μ_n) siguen la formulación "
        "de Sturm-Liouville Caótico."
    ))

    # ── 3. Componentes ──
    add_heading(doc, "3. Componentes del Hexágono", 1)

    add_heading(doc, "3.1 Conexión Hexagonal", 2)
    add_para(doc, (
        "Cada componente del hexágono se conecta al JEPA Hub mediante proyecciones "
        "bidireccionales que transforman entre el espacio de trabajo del componente "
        "y el espacio latente JEPA:"
    ))
    add_code(doc, """class HexagonalTorus(nn.Module):
    def __init__(self, latent_dim, component_dim):
        super().__init__()
        self.to_latent = nn.Linear(component_dim, latent_dim, bias=False)
        self.from_latent = nn.Linear(latent_dim, component_dim, bias=False)""")

    add_heading(doc, "3.2 Transformer (Flash Attention 2)", 2)
    add_para(doc, (
        "10 bloques transformer con Flash Attention 2, RoPE (Rotary Position Embeddings), "
        "GQA (Grouped Query Attention, n_kv=4), y NTK-aware scaling. "
        "Cada bloque procesa desde el espacio JEPA y retorna a JEPA."
    ))

    add_heading(doc, "3.3 LIF-SNN (Spiking Neural Network)", 2)
    add_para(doc, (
        "Neuronas Leaky Integrate-and-Fire con threshold dinámico (0.001) y tau=0.8. "
        "El SNN dispara al ~47% de tasa de spike, contribuyendo procesamiento temporal "
        "al espacio latente JEPA."
    ))

    add_heading(doc, "3.4 SSM (State Space Model)", 2)
    add_para(doc, (
        "Modelo de espacio de estado con state_dim=64, que captura dependencias de largo "
        "alcance en el espacio latente JEPA."
    ))

    add_heading(doc, "3.5 Synapsis (Memoria Persistente)", 2)
    add_para(doc, (
        "Memoria toroidal con 256 slots y top-5 retrieval. Solo activa en inferencia "
        "(desactivada durante entrenamiento para evitar el efecto de repetición 'the the the')."
    ))

    # ── 4. Flujo Toroidal ──
    add_heading(doc, "4. Flujo Toroidal", 1)
    add_para(doc, (
        "El forward pass ejecuta N ciclos alrededor del toroide hexagonal. "
        "En cada ciclo, todos los componentes leen del JEPA Hub, procesan, y "
        "retornan al JEPA Hub para integración."
    ))
    add_code(doc, """FASE 1: EMBEDDING → JEPA HUB
  h = tok_emb(x)
  h = HSAQ(h, 5%)
  latent = jepa_enc(emb_to_jepa(h))
  latent = HSAQ(latent, 5%)

FASE 2: CICLOS TOROIDALES (×N)
  for cycle in range(n_cycles):
    # Transformer ← JEPA → Transformer
    t_out = transformer(t_from_jepa(latent))
    t_latent = t_to_jepa(HSAQ(t_out, 8-15%))

    # SNN ← JEPA → SNN
    s_out = snn(s_from_jepa(latent))
    s_latent = s_to_jepa(HSAQ(s_out, 10%))

    # SSM ← JEPA → SSM
    ssm_out = ssm(ssm_from_jepa(latent))
    ssm_latent = ssm_to_jepa(HSAQ(ssm_out, 5%))

    # JEPA INTEGRA
    latent = jepa_enc((latent + t_latent + s_latent + ssm_latent) / 4)
    latent = HSAQ(latent, 5%)

FASE 3: JEPA PREDICE
  jepa_mse = MSE(jepa_pred(latent[:-1]), latent[1:].detach())

FASE 4: HEAD
  out = norm(latent)
  logits = head(out)""")

    # ── 5. HSAQ ──
    add_heading(doc, "5. HSAQ en el Toroide", 1)
    add_para(doc, (
        "HSAQ se aplica en todas las aristas del hexágono con sparsity calibrada "
        "por componente. Para 2 ciclos toroidales, hay 22 puntos de aplicación HSAQ:"
    ))
    make_table(doc,
        ["Ciclo", "Arista", "Sparsity", "Target"],
        [
            ["Inicio", "Embedding", "5%", "Preservar entrada"],
            ["", "JEPA Hub", "5%", "Preservar latente"],
            ["Ciclo 1", "Transformer t2/t5/t8", "8/12/15%", "Progresivo"],
            ["", "SNN", "10%", "Spikes controlados"],
            ["", "SSM", "5%", "Estado latente"],
            ["", "JEPA integración", "5%", "Preservar fusión"],
            ["Ciclo 2", "Transformer t2/t5/t8", "8/12/15%", "Progresivo"],
            ["", "SNN", "10%", "Spikes controlados"],
            ["", "SSM", "5%", "Estado latente"],
            ["", "JEPA integración", "5%", "Preservar fusión"],
        ]
    )
    add_caption(doc, "Tabla 1: Puntos HSAQ en el toroide hexagonal (2 ciclos)")

    # ── 6. Parámetros ──
    add_heading(doc, "6. Parámetros del Modelo", 1)
    make_table(doc,
        ["Parámetro", "Valor"],
        [
            ["Parámetros totales", "140.9M"],
            ["Dimensión (dim)", "896"],
            ["Capas transformer", "10"],
            ["Cabezas de atención", "8 query / 4 KV (GQA)"],
            ["Dimensión latente JEPA", "896"],
            ["Dimensión SNN", "448"],
            ["Estado SSM", "64"],
            ["Ciclos toroidales (n_cycles)", "2"],
            ["Vocabulario", "1024 (char-level)"],
            ["Optimizer", "SGD Nesterov (momentum=0.9)"],
            ["Learning rate", "5×10⁻⁴ (cosine decay)"],
            ["HSAQ sparsity base", "0.3 (escalonada por capa)"],
        ]
    )
    add_caption(doc, "Tabla 2: Parámetros de M.A.T.E.R.I.A. V4")

    # ── 7. Resultados ──
    add_heading(doc, "7. Resultados de Entrenamiento", 1)
    make_table(doc,
        ["Métrica", "V3 (lineal)", "V4 (toroidal)", "Mejora"],
        [
            ["Accuracy", "19.5%", "28.9%+", "+48%"],
            ["Perplexity", "107", "83.8", "-22%"],
            ["Loss", "3.73", "3.88", "Equivalente"],
            ["SNN spike rate", "0.0 (inactivo)", "0.47 (47%)", "✅"],
            ["Info preservada HSAQ", "0.7%", "~52%", "75×"],
            ["OOMs", "Frecuentes", "0", "✅"],
        ]
    )
    add_caption(doc, "Tabla 3: Comparativa V3 lineal vs V4 toroidal")

    # ── 8. Referencias ──
    add_heading(doc, "8. Referencias", 1)
    add_bullet(doc, "HSAQ Standard v1.1 — docs/HSAQ_STANDARD.docx")
    add_bullet(doc, "Paper Científico M.A.T.E.R.I.A. V4 — Documentos IA LATAM")
    add_bullet(doc, "K = √(π·e·γ) = 2.781042 — Constante de acoplamiento espectral")
    add_bullet(doc, "Flash Attention 2 — Dao et al. 2023")
    add_bullet(doc, "RoPE — Su et al. RoFormer 2023")

    # ── Cierre ──
    add_para(doc, "", space_after=12)
    add_para(doc, "— Fin del documento —", italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(path)
    print(f"✅ MATERIA DOCX: {path}")


# ═══════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    generate_hsaq_docx()
    generate_materia_docx()
    print("\n✅ Ambos documentos DOCX generados correctamente")
